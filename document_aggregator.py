# -*- coding: utf-8 -*-
"""
DocumentAggregator v3.0 – Công cụ Tổng hợp & Phân tích Tài liệu Đa năng
=========================================================================
Dự án: DU AN SUC KHOE MR PHI
Hỗ trợ: PDF (Scan & Text), Word (.docx), Excel (.xlsx)
Tính năng:
  - Tự động nhận diện & phân loại: Bảng số liệu (Structured) vs Văn bản (Text/Legal)
  - Tích hợp OCR siêu tốc (RapidOCR) cho PDF scan có bộ đệm lưu (cache)
  - Tự động trích xuất thông tin then chốt: Số tiền, Ngày tháng, Số hợp đồng/VB, Các bên
  - Xuất Excel chuyên nghiệp nhiều Sheet có định dạng đẹp mắt
  - Tổng hợp thống kê tự động từ dữ liệu đã quét (không hardcode)
"""

import os
import re
import io
import json
import hashlib
import warnings
import pandas as pd
import pdfplumber
import fitz          # PyMuPDF
from docx import Document
from pathlib import Path

warnings.filterwarnings("ignore", category=UserWarning)

# Khởi tạo OCR engine (singleton)
try:
    from rapidocr_onnxruntime import RapidOCR
    OCR_ENGINE = RapidOCR()
except Exception:
    OCR_ENGINE = None

# Thư mục cache OCR để tăng tốc độ tải
CACHE_DIR = Path(os.getenv("TEMP", "/tmp")) / "doc_aggregator_ocr_cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)


class DocumentAggregator:
    """
    Hệ thống Tổng hợp & Phân tích Tài liệu Đa định dạng v3.0
    Phiên bản đa năng – không gắn cứng vào dự án cụ thể.
    """

    SUPPORTED_EXTS = {".pdf", ".docx", ".doc", ".xlsx", ".xls"}

    def __init__(self, project_name: str = "Dự Án Sức Khỏe Mr Phi", verbose: bool = True):
        self.project_name = project_name
        self.verbose = verbose
        # Dữ liệu phân loại
        self.excel_tables = []       # Danh sách DataFrame từ Excel
        self.doc_tables   = []       # Danh sách bảng biểu từ Word / PDF
        self.text_docs    = []       # Danh sách văn bản (Word, PDF text / OCR)
        self.key_insights = []       # Thông tin then chốt trích xuất (tiền, ngày, số VB...)
        self.errors       = []       # Danh sách lỗi
        self.files_read   = []       # Danh sách file đã đọc
        self._workspace   = None     # Thư mục workspace đã quét

    # ─────────────────────────────────────────────────────────────
    # TRÍCH XUẤT THÔNG TIN THEN CHỐT (REGEX & PATTERN MATCHING)
    # ─────────────────────────────────────────────────────────────
    @staticmethod
    def extract_key_entities(text: str, file_name: str) -> dict:
        """Trích xuất số tiền, ngày tháng, số hợp đồng, các bên từ văn bản."""
        if not text:
            return {}

        # 1. Tìm các khoản tiền lớn
        money_patterns = [
            r'(\d{1,3}(?:\.\d{3}){2,4})\s*(?:đồng|VNĐ|vnd|VND|đ)?',  # 140.087.431.751
            r'(\d+(?:[.,]\d+)?)\s*(?:tỷ|triệu)\s*(?:đồng|VNĐ|VND)?', # 80 tỷ đồng
        ]
        moneys = []
        for pat in money_patterns:
            matches = re.findall(pat, text, re.IGNORECASE)
            for m in matches:
                if len(m) > 4 or 'tỷ' in pat:
                    moneys.append(m.strip())

        # 2. Tìm ngày tháng
        dates = re.findall(
            r'(?:ngày\s+)?(\d{1,2}[/-]\d{1,2}[/-]\d{4}|\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})',
            text, re.IGNORECASE
        )

        # 3. Tìm số hợp đồng / công văn
        contract_nums = re.findall(
            r'(?:Số|HĐ số|Hợp đồng số|Quyết định số|QĐ số|CV số|Công văn số)\s*[:.]\s*([0-9A-Za-z/_\-]+(?:\/[0-9A-Za-z/_\-]+)*)',
            text, re.IGNORECASE
        )

        # 4. Tìm các bên liên quan (tổng quát – tìm tên công ty/tổ chức)
        parties = []
        # Tìm "Công ty ..." patterns
        company_patterns = re.findall(
            r'(?:Công ty\s+(?:TNHH|Cổ phần|CP|TNHH MTV)\s+[^\n.,;]{3,50})',
            text, re.IGNORECASE
        )
        for cp in company_patterns[:5]:
            clean = cp.strip()
            if clean and clean not in parties:
                parties.append(clean)

        # Tìm ngân hàng
        bank_patterns = re.findall(
            r'(?:Ngân hàng|NH|Bank)\s+[^\n.,;]{3,40}',
            text, re.IGNORECASE
        )
        for bp in bank_patterns[:3]:
            clean = bp.strip()
            if clean and clean not in parties:
                parties.append(clean)

        # Tìm tên riêng viết hoa (Ông/Bà + Tên)
        person_patterns = re.findall(
            r'(?:Ông|Bà|Mr\.|Mrs\.)\s+([A-ZÀ-Ỹ][a-zà-ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹ]+){1,4})',
            text
        )
        for pp in person_patterns[:4]:
            name = pp.strip()
            if name and len(name) > 3 and name not in parties:
                parties.append(name)

        return {
            "file_name": file_name,
            "moneys": list(dict.fromkeys(moneys))[:8],
            "dates": list(dict.fromkeys(dates))[:5],
            "contract_nums": list(dict.fromkeys(contract_nums))[:3],
            "parties": list(dict.fromkeys(parties)),
        }

    # ─────────────────────────────────────────────────────────────
    # ĐỌC PDF (TEXT + BẢNG + OCR SCAN CÓ CACHE)
    # ─────────────────────────────────────────────────────────────
    def read_pdf(self, filepath: str) -> dict:
        """Đọc file PDF, tự động OCR nếu là PDF scan."""
        path = Path(filepath)
        records_tables = []
        records_text   = []

        # Tạo hash file để cache OCR
        file_hash = hashlib.md5(path.read_bytes()[:1024*1024]).hexdigest()
        cache_file = CACHE_DIR / f"{path.stem}_{file_hash}.json"

        # Nếu đã có cache OCR
        if cache_file.exists():
            try:
                with open(cache_file, "r", encoding="utf-8") as cf:
                    cached_data = json.load(cf)
                    for row in cached_data:
                        records_text.append(row)
                return {"tables": records_tables, "text": records_text}
            except Exception:
                pass

        try:
            with pdfplumber.open(filepath) as pdf:
                doc_mu = fitz.open(filepath)

                for page_idx, page in enumerate(pdf.pages):
                    page_num = page_idx + 1
                    # 1. Trích xuất bảng
                    tables = page.extract_tables()
                    if tables:
                        for tbl_idx, tbl in enumerate(tables):
                            if tbl and len(tbl) > 1:
                                header = [str(h).strip() if h else f"col_{c}" for c, h in enumerate(tbl[0])]
                                df_t = pd.DataFrame(tbl[1:], columns=header)
                                df_t["__source_file__"] = path.name
                                df_t["__page__"]        = page_num
                                records_tables.append(df_t)

                    # 2. Trích xuất văn bản (pdfplumber -> PyMuPDF -> RapidOCR)
                    text = (page.extract_text() or "").strip()
                    ocr_used = False

                    if not text:
                        text = doc_mu[page_idx].get_text().strip()

                    # Nếu là trang scan (không có text) -> Chạy OCR
                    if not text and OCR_ENGINE is not None:
                        pix = doc_mu[page_idx].get_pixmap(dpi=130)
                        ocr_res, _ = OCR_ENGINE(pix.tobytes("png"))
                        if ocr_res:
                            text = "\n".join([line[1] for line in ocr_res])
                            ocr_used = True

                    if text:
                        records_text.append({
                            "file_name": path.name,
                            "page": page_num,
                            "type": "OCR_Scan" if ocr_used else "Digital_PDF",
                            "content": text,
                            "char_count": len(text),
                        })

                doc_mu.close()

            # Lưu cache kết quả text
            if records_text:
                try:
                    with open(cache_file, "w", encoding="utf-8") as cf:
                        json.dump(records_text, cf, ensure_ascii=False, indent=2)
                except Exception:
                    pass

        except Exception as e:
            if self.verbose:
                print(f"  [!] Lỗi đọc PDF {path.name}: {e}")
            self.errors.append({"file": path.name, "error": str(e)})

        return {"tables": records_tables, "text": records_text}

    # ─────────────────────────────────────────────────────────────
    # ĐỌC WORD (.DOCX)
    # ─────────────────────────────────────────────────────────────
    def read_word(self, filepath: str) -> dict:
        """Đọc file Word (.docx) - tách rõ Bảng biểu và Đoạn văn."""
        path = Path(filepath)
        records_tables = []
        records_text   = []

        try:
            doc = Document(filepath)

            # 1. Đọc các bảng
            for tbl_idx, tbl in enumerate(doc.tables):
                data = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
                if data and len(data) > 1:
                    max_cols = max(len(r) for r in data)
                    data = [r + [""] * (max_cols - len(r)) for r in data]
                    seen, header = {}, []
                    for h in data[0]:
                        h = h or "Cột"
                        seen[h] = seen.get(h, 0) + 1
                        header.append(f"{h}_{seen[h]}" if seen[h] > 1 else h)
                    df_t = pd.DataFrame(data[1:], columns=header)
                    df_t["__source_file__"] = path.name
                    df_t["__table__"]       = tbl_idx + 1
                    records_tables.append(df_t)

            # 2. Đọc toàn bộ đoạn văn
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if paras:
                full_text = "\n\n".join(paras)
                records_text.append({
                    "file_name": path.name,
                    "page": 1,
                    "type": "Word_Document",
                    "content": full_text,
                    "char_count": len(full_text),
                })

        except Exception as e:
            if self.verbose:
                print(f"  [!] Lỗi đọc Word {path.name}: {e}")
            self.errors.append({"file": path.name, "error": str(e)})

        return {"tables": records_tables, "text": records_text}

    # ─────────────────────────────────────────────────────────────
    # ĐỌC EXCEL (.XLSX / .XLS)
    # ─────────────────────────────────────────────────────────────
    def read_excel(self, filepath: str) -> list[pd.DataFrame]:
        """Đọc toàn bộ các sheet của file Excel."""
        path = Path(filepath)
        dfs = []
        try:
            xl = pd.ExcelFile(filepath)
            for sheet in xl.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet, dtype=str)
                df = df.dropna(how="all")
                if not df.empty:
                    df["__file_name__"] = path.name
                    df["__sheet__"]     = sheet
                    dfs.append(df)
        except Exception as e:
            if self.verbose:
                print(f"  [!] Lỗi đọc Excel {path.name}: {e}")
            self.errors.append({"file": path.name, "error": str(e)})
        return dfs

    # ─────────────────────────────────────────────────────────────
    # ĐỌC FILE ĐƠN LẺ
    # ─────────────────────────────────────────────────────────────
    def read_file(self, filepath: str) -> dict:
        path = Path(filepath)
        if not path.exists():
            raise FileNotFoundError(f"File không tồn tại: {filepath}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTS:
            raise ValueError(f"Định dạng {ext} chưa được hỗ trợ.")

        res = {"type": ext, "tables": [], "text": [], "excel": []}

        if ext == ".pdf":
            out = self.read_pdf(str(path))
            res["tables"] = out["tables"]
            res["text"]   = out["text"]
            self.doc_tables.extend(out["tables"])
            self.text_docs.extend(out["text"])

        elif ext in (".docx", ".doc"):
            out = self.read_word(str(path))
            res["tables"] = out["tables"]
            res["text"]   = out["text"]
            self.doc_tables.extend(out["tables"])
            self.text_docs.extend(out["text"])

        elif ext in (".xlsx", ".xls"):
            dfs = self.read_excel(str(path))
            res["excel"] = dfs
            self.excel_tables.extend(dfs)

        # Trích xuất thông tin then chốt từ văn bản
        for item in res["text"]:
            insights = self.extract_key_entities(item.get("content", ""), path.name)
            if insights:
                self.key_insights.append(insights)

        self.files_read.append(path.name)
        return res

    # ─────────────────────────────────────────────────────────────
    # ĐỌC CẢ THƯ MỤC
    # ─────────────────────────────────────────────────────────────
    def read_folder(self, folder_path: str, recursive: bool = True):
        folder = Path(folder_path)
        if not folder.is_dir():
            raise NotADirectoryError(f"Không tìm thấy folder: {folder_path}")

        self._workspace = str(folder)

        glob_fn = folder.rglob if recursive else folder.glob
        all_files = [f for f in sorted(glob_fn("*"))
                     if f.is_file() and not f.name.startswith("~$")
                     and f.suffix.lower() in self.SUPPORTED_EXTS]

        if self.verbose:
            print(f"\n📂 Đang quét {len(all_files)} file trong '{folder.name}'...")

        for f in all_files:
            if self.verbose:
                print(f"  ▶ Đang đọc: {f.name} ...", end=" ", flush=True)
            self.read_file(str(f))
            if self.verbose:
                print("✅")

        return self

    # ─────────────────────────────────────────────────────────────
    # TỔNG HỢP THỐNG KÊ TỰ ĐỘNG TỪ DỮ LIỆU ĐÃ QUÉT
    # ─────────────────────────────────────────────────────────────
    def get_project_summary(self) -> dict:
        """
        Tự động tổng hợp thống kê từ dữ liệu đã quét.
        Trả về cấu trúc dict chung (không hardcode số liệu).
        """
        # Thu thập tất cả số tiền tìm được
        all_moneys = []
        all_dates  = []
        all_parties = []
        for ins in self.key_insights:
            all_moneys.extend(ins.get("moneys", []))
            all_dates.extend(ins.get("dates", []))
            all_parties.extend(ins.get("parties", []))

        unique_files = list(dict.fromkeys(self.files_read))

        return {
            "tong_quan": {
                "du_an": self.project_name,
                "workspace": self._workspace or "Chưa quét",
                "tong_so_file": len(unique_files),
                "tong_trang_van_ban": len(self.text_docs),
                "tong_sheet_excel": len(self.excel_tables),
                "tong_bang_bieu": len(self.doc_tables),
                "tong_loi": len(self.errors),
                "so_tien_phat_hien": list(dict.fromkeys(all_moneys))[:15],
                "ngay_thang_phat_hien": list(dict.fromkeys(all_dates))[:10],
                "cac_ben_lien_quan": list(dict.fromkeys(all_parties))[:10],
            },
            "danh_sach_file": unique_files,
        }

    # ─────────────────────────────────────────────────────────────
    # XUẤT FILE EXCEL CHUYÊN NGHIỆP (CLEAN MULTI-SHEET)
    # ─────────────────────────────────────────────────────────────
    def export_excel(self, output_path: str = "Ket_qua_tong_hop.xlsx") -> str:
        """Xuất toàn bộ dữ liệu ra Excel chuẩn định dạng nhiều sheet."""
        summary = self.get_project_summary()

        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            # Sheet 1: Tổng quan dự án
            rows_tq = [
                {"Chỉ tiêu": "Dự án", "Nội dung / Giá trị": summary["tong_quan"]["du_an"]},
                {"Chỉ tiêu": "Thư mục dữ liệu", "Nội dung / Giá trị": summary["tong_quan"]["workspace"]},
                {"Chỉ tiêu": "Tổng số file đã quét", "Nội dung / Giá trị": str(summary["tong_quan"]["tong_so_file"])},
                {"Chỉ tiêu": "Tổng trang văn bản (Word/PDF/OCR)", "Nội dung / Giá trị": str(summary["tong_quan"]["tong_trang_van_ban"])},
                {"Chỉ tiêu": "Tổng sheet Excel", "Nội dung / Giá trị": str(summary["tong_quan"]["tong_sheet_excel"])},
                {"Chỉ tiêu": "Tổng bảng biểu trong tài liệu", "Nội dung / Giá trị": str(summary["tong_quan"]["tong_bang_bieu"])},
                {"Chỉ tiêu": "Số tiền phát hiện trong hồ sơ", "Nội dung / Giá trị": " | ".join(summary["tong_quan"]["so_tien_phat_hien"][:8])},
                {"Chỉ tiêu": "Ngày tháng phát hiện", "Nội dung / Giá trị": " | ".join(summary["tong_quan"]["ngay_thang_phat_hien"][:6])},
                {"Chỉ tiêu": "Các bên liên quan", "Nội dung / Giá trị": " | ".join(summary["tong_quan"]["cac_ben_lien_quan"][:6])},
            ]
            pd.DataFrame(rows_tq).to_excel(writer, sheet_name="Tong_Quan_Du_An", index=False)

            # Sheet 2: Danh sách file đã quét
            if summary["danh_sach_file"]:
                df_files = pd.DataFrame([
                    {"STT": idx + 1, "Tên file": fname}
                    for idx, fname in enumerate(summary["danh_sach_file"])
                ])
                df_files.to_excel(writer, sheet_name="Danh_Sach_File", index=False)

            # Sheet 3: Dữ liệu bảng tính Excel (nếu có)
            if self.excel_tables:
                df_excel_all = pd.concat(self.excel_tables, ignore_index=True)
                df_excel_all.to_excel(writer, sheet_name="Du_Lieu_Bang_Tinh", index=False)

            # Sheet 4: Trích yếu các văn bản & Hợp đồng (Insights)
            if self.key_insights:
                insights_rows = []
                for item in self.key_insights:
                    insights_rows.append({
                        "Tên file": item.get("file_name", ""),
                        "Số tiền phát hiện": ", ".join(item.get("moneys", [])),
                        "Ngày tháng phát hiện": ", ".join(item.get("dates", [])),
                        "Số HĐ / Công văn": ", ".join(item.get("contract_nums", [])),
                        "Các bên liên quan": ", ".join(item.get("parties", [])),
                    })
                df_ins = pd.DataFrame(insights_rows).drop_duplicates(subset=["Tên file"])
                df_ins.to_excel(writer, sheet_name="Trich_Yeu_Phap_Ly", index=False)

            # Sheet 5: Chi tiết nội dung văn bản (Word & PDF OCR)
            if self.text_docs:
                df_docs = pd.DataFrame(self.text_docs).rename(columns={
                    "file_name": "Tên file",
                    "page": "Trang",
                    "type": "Loại tài liệu",
                    "char_count": "Số ký tự",
                    "content": "Nội dung trích xuất",
                })
                df_docs.to_excel(writer, sheet_name="Chi_Tiet_Noi_Dung_Van_Ban", index=False)

            # Sheet 6: Danh mục lỗi (nếu có)
            if self.errors:
                pd.DataFrame(self.errors).to_excel(writer, sheet_name="Nhat_Ky_Loi", index=False)

        if self.verbose:
            print(f"\n✅ Đã xuất file Excel chuẩn: {output_path}")

        return output_path
