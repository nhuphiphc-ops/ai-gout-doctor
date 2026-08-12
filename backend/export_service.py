import io
from datetime import date
from typing import List
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import models
import ai_engine

def export_to_excel(logs: List[models.HealthLog]) -> io.BytesIO:
    """
    Exports health logs to an Excel spreadsheet.
    """
    data = []
    for log in logs:
        foods_str = ", ".join([f.food_name for f in log.foods])
        meds_str = ", ".join([m.med_name for m in log.medications])
        
        # Add high-purine flags as list
        purines = []
        if log.had_beer: purines.append("Bia")
        if log.had_alcohol: purines.append("Rượu/Cồn")
        if log.had_seafood: purines.append("Hải sản")
        if log.had_organ_meat: purines.append("Nội tạng")
        if log.had_red_meat: purines.append("Thịt đỏ")
        if log.had_sweets: purines.append("Đồ ngọt")
        purines_str = ", ".join(purines)

        data.append({
            "Ngày": log.log_date,
            "Cân nặng (kg)": log.weight,
            "Huyết áp Systolic": log.bp_systolic,
            "Huyết áp Diastolic": log.bp_diastolic,
            "Nhịp tim (bpm)": log.heart_rate,
            "Thời gian ngủ (giờ)": log.sleep_duration,
            "Chất lượng ngủ (1-10)": log.sleep_quality,
            "Đau khớp chân": "Có" if log.joint_pain else "Không",
            "Mức độ đau (0-10)": log.pain_severity,
            "Số bước chân": log.steps,
            "Nước uống (lít)": log.water_intake,
            "Thực phẩm purin ăn vào": purines_str,
            "Món ăn chi tiết": foods_str,
            "Thuốc đã uống": meds_str,
            "Điểm Gout": log.gout_score,
            "Điểm Tim mạch": log.cardio_score,
            "Điểm Chuyển hóa": log.metabolic_score,
            "Điểm QoL": log.qol_score
        })

    df = pd.DataFrame(data)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Nhat_Ky_Suc_Khoe')
    
    output.seek(0)
    return output

def export_to_docx(logs: List[models.HealthLog], user: models.User, correlations: List[dict]) -> io.BytesIO:
    """
    Exports health summary and analysis to a Word document (.docx).
    """
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    r = title.add_run("BÁO CÁO PHÂN TÍCH SỨC KHỎE CÁ NHÂN")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
    title.alignment = 1 # Center
    
    doc.add_paragraph(f"Ngày lập báo cáo: {date.today().strftime('%d/%m/%Y')}")
    
    # User Profile
    doc.add_heading("1. Thông tin cá nhân", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Họ và tên: ").bold = True
    p.add_run(f"{user.name or 'Người dùng'}\n")
    p.add_run(f"Tuổi: ").bold = True
    p.add_run(f"{user.age} tuổi\n")
    p.add_run(f"Chiều cao: ").bold = True
    p.add_run(f"{user.height:.2f} m\n")
    p.add_run(f"Cân nặng hiện tại (trung bình): ").bold = True
    avg_weight = sum([l.weight for l in logs if l.weight]) / len([l for l in logs if l.weight]) if logs else 62.5
    p.add_run(f"{avg_weight:.1f} kg\n")
    p.add_run(f"Chỉ số BMI: ").bold = True
    p.add_run(f"{avg_weight / (user.height ** 2):.1f}\n")
    p.add_run(f"Tiền sử bệnh án: ").bold = True
    p.add_run("Bị gout từ 11 năm trước, thỉnh thoảng đau nhức chân nhẹ.")

    # Health averages
    doc.add_heading("2. Chỉ số sức khỏe trung bình", level=1)
    
    # Table of averages
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ số'
    hdr_cells[1].text = 'Giá trị trung bình'
    hdr_cells[2].text = 'Mục tiêu'
    hdr_cells[3].text = 'Đánh giá'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    # Compute averages
    count = len(logs)
    if count > 0:
        avg_bp_sys = int(sum([l.bp_systolic for l in logs if l.bp_systolic]) / len([l for l in logs if l.bp_systolic])) if [l for l in logs if l.bp_systolic] else 120
        avg_bp_dia = int(sum([l.bp_diastolic for l in logs if l.bp_diastolic]) / len([l for l in logs if l.bp_diastolic])) if [l for l in logs if l.bp_diastolic] else 80
        avg_steps = int(sum([l.steps for l in logs]) / count)
        avg_water = sum([l.water_intake for l in logs]) / count
        avg_sleep = sum([l.sleep_duration for l in logs if l.sleep_duration]) / len([l for l in logs if l.sleep_duration]) if [l for l in logs if l.sleep_duration] else 7.0
        pain_days = sum([1 for l in logs if l.joint_pain])
        
        metrics = [
            ("Huyết áp (mmHg)", f"{avg_bp_sys}/{avg_bp_dia}", "< 120/80", "Bình thường" if avg_bp_sys < 130 else "Cần theo dõi"),
            ("Số bước chân", f"{avg_steps:,}", ">= 7,000", "Tốt" if avg_steps >= 7000 else "Cần tăng vận động"),
            ("Lượng nước uống (L)", f"{avg_water:.1f} L", "2.5 L", "Đủ nước" if avg_water >= 2.3 else "Thiếu nước nhẹ"),
            ("Thời gian ngủ", f"{avg_sleep:.1f} giờ", "7.0 - 8.0", "Đủ giấc" if 7 <= avg_sleep <= 8.5 else "Chưa tối ưu"),
            ("Số ngày đau khớp", f"{pain_days} ngày / {count} ngày", "0 ngày", "Cảnh báo bùng phát" if pain_days > 0 else "Ổn định")
        ]
        
        for name, avg, target, eval_str in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = avg
            row_cells[2].text = target
            row_cells[3].text = eval_str

    # Food correlation
    doc.add_heading("3. Báo cáo Tương quan Thực phẩm & Cơn đau khớp (30 ngày)", level=1)
    doc.add_paragraph("Dưới đây là danh sách thực phẩm được ghi nhận tiêu thụ nhiều nhất trước khi xuất hiện cảm giác đau nhức chân:")
    
    if correlations:
        corr_table = doc.add_table(rows=1, cols=4)
        corr_hdr = corr_table.rows[0].cells
        corr_hdr[0].text = 'Tên thực phẩm/Thói quen'
        corr_hdr[1].text = 'Số lần ăn trước cơn đau'
        corr_hdr[2].text = 'Tổng số lần tiêu thụ'
        corr_hdr[3].text = 'Tỷ lệ kích hoạt đau (%)'
        
        for cell in corr_hdr:
            cell.paragraphs[0].runs[0].font.bold = True
            
        for item in correlations[:20]:
            row = corr_table.add_row().cells
            row[0].text = str(item["food_name"])
            row[1].text = str(item["pain_incidents_with_food"])
            row[2].text = str(item["total_consumption"])
            row[3].text = f"{item['correlation_percentage']}%"
    else:
        doc.add_paragraph("Chưa có đủ dữ liệu tương quan cơn đau (Yêu cầu ít nhất 1 đợt đau chân được báo cáo).")

    # Recommendations
    doc.add_heading("4. Khuyến nghị từ Bác sĩ gia đình AI", level=1)
    
    # Standard recommendations
    p_rec = doc.add_paragraph()
    p_rec.add_run("Dinh dưỡng:\n").bold = True
    p_rec.add_run("- Tăng cường: Cá sông thịt trắng, rau bắp cải, súp lơ, cần tây, anh đào (cherries), chanh.\n")
    p_rec.add_run("- Hạn chế tối đa: Hải sản có vỏ, thịt bò/trâu, lòng mề động vật, nước ngọt đóng chai, bia rượu.\n\n")
    p_rec.add_run("Lối sống và Vận động:\n").bold = True
    p_rec.add_run("- Duy trì đi bộ tối thiểu 6000-8000 bước mỗi ngày khi khớp không đau.\n")
    p_rec.add_run("- Uống đủ 2.5 lít nước ấm rải đều trong ngày (không uống dồn dập).\n")
    p_rec.add_run("- Kéo giãn cổ chân và gân kheo mỗi tối trước khi ngủ để tăng tuần hoàn dịch khớp.")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_to_pdf(logs: List[models.HealthLog], user: models.User, correlations: List[dict]) -> io.BytesIO:
    """
    Exports health summary and analysis to a beautiful PDF document using reportlab.
    """
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        textColor=colors.HexColor('#10b981'),
        spaceAfter=15,
        alignment=1
    )
    
    h1_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=colors.HexColor('#1f2937'),
        spaceBefore=12,
        spaceAfter=6
    )
    
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor('#374151'),
        spaceAfter=8
    )
    
    story = []
    
    # Title
    story.append(Paragraph("BÁO CÁO SỨC KHỎE TOÀN DIỆN & NGUY CƠ GOUT", title_style))
    story.append(Paragraph(f"Thời gian lập báo cáo: {date.today().strftime('%d/%m/%Y')}", body_style))
    story.append(Spacer(1, 15))
    
    # 1. Profile
    story.append(Paragraph("1. Thông tin chỉ số cá nhân", h1_style))
    profile_text = f"""
    <b>Họ và tên:</b> {user.name or 'Người dùng'} &nbsp;&nbsp;&nbsp;&nbsp; <b>Tuổi:</b> {user.age} tuổi<br/>
    <b>Chiều cao:</b> {user.height:.2f} m &nbsp;&nbsp;&nbsp;&nbsp; <b>Cân nặng mục tiêu:</b> {user.target_weight} kg<br/>
    <b>Tiền sử bệnh án:</b> Đã chẩn đoán Gout từ 11 năm trước. Sức khỏe khớp hiện ổn định.
    """
    story.append(Paragraph(profile_text, body_style))
    story.append(Spacer(1, 10))
    
    # 2. Health averages
    story.append(Paragraph("2. Chỉ số sức khỏe trung bình", h1_style))
    
    # Build data table
    count = len(logs)
    if count > 0:
        avg_bp_sys = int(sum([l.bp_systolic for l in logs if l.bp_systolic]) / len([l for l in logs if l.bp_systolic])) if [l for l in logs if l.bp_systolic] else 120
        avg_bp_dia = int(sum([l.bp_diastolic for l in logs if l.bp_diastolic]) / len([l for l in logs if l.bp_diastolic])) if [l for l in logs if l.bp_diastolic] else 80
        avg_steps = int(sum([l.steps for l in logs]) / count)
        avg_water = sum([l.water_intake for l in logs]) / count
        avg_sleep = sum([l.sleep_duration for l in logs if l.sleep_duration]) / len([l for l in logs if l.sleep_duration]) if [l for l in logs if l.sleep_duration] else 7.0
        pain_days = sum([1 for l in logs if l.joint_pain])

        table_data = [
            [Paragraph("<b>Chỉ số sức khỏe</b>", body_style), 
             Paragraph("<b>Trung bình thực tế</b>", body_style), 
             Paragraph("<b>Mục tiêu y khoa</b>", body_style), 
             Paragraph("<b>Đánh giá sức khỏe</b>", body_style)],
            [Paragraph("Huyết áp tĩnh", body_style), f"{avg_bp_sys}/{avg_bp_dia} mmHg", "< 120/80", "Tốt" if avg_bp_sys < 130 else "Hơi cao"],
            [Paragraph("Vận động (Bước chân)", body_style), f"{avg_steps:,}", ">= 7,000 bước", "Tốt" if avg_steps >= 7000 else "Ít vận động"],
            [Paragraph("Uống nước hàng ngày", body_style), f"{avg_water:.1f} lít", "2.5 lít", "Đạt chuẩn" if avg_water >= 2.3 else "Thiếu nước"],
            [Paragraph("Giấc ngủ", body_style), f"{avg_sleep:.1f} giờ", "7.0 - 8.0 giờ", "Đủ giấc" if 7 <= avg_sleep <= 8.5 else "Chưa sâu giấc"],
            [Paragraph("Tần suất đau nhức khớp", body_style), f"{pain_days} ngày / {count} ngày", "0 ngày", "Bình thường" if pain_days == 0 else "Có dấu hiệu bùng phát"]
        ]
        
        t = Table(table_data, colWidths=[150, 110, 100, 150])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e2e8f0')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t)
    else:
        story.append(Paragraph("Chưa ghi nhận dữ liệu nhật ký sức khỏe.", body_style))
        
    story.append(Spacer(1, 15))
    
    # 3. Correlations
    story.append(Paragraph("3. Bản đồ tương quan Thực phẩm - Đau khớp (30 ngày)", h1_style))
    story.append(Paragraph("Danh sách 10 thực phẩm có xác suất liên đới cao nhất tới các đợt đau nhức chân của anh:", body_style))
    
    if correlations:
        corr_table_data = [
            [Paragraph("<b>Thực phẩm / Thói quen</b>", body_style), 
             Paragraph("<b>Ăn trước đau (lần)</b>", body_style), 
             Paragraph("<b>Tổng ngày tiêu thụ</b>", body_style), 
             Paragraph("<b>Xác suất kích hoạt đau</b>", body_style)]
        ]
        for item in correlations[:10]:
            pct = item["correlation_percentage"]
            color_hex = '#ef4444' if pct >= 60 else ('#f59e0b' if pct >= 30 else '#10b981')
            corr_table_data.append([
                Paragraph(item["food_name"], body_style),
                str(item["pain_incidents_with_food"]),
                str(item["total_consumption"]),
                Paragraph(f"<font color='{color_hex}'><b>{pct}%</b></font>", body_style)
            ])
            
        ct = Table(corr_table_data, colWidths=[180, 110, 110, 110])
        ct.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e2e8f0')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('PADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(ct)
    else:
        story.append(Paragraph("Chưa ghi nhận đủ số ngày đau khớp để tính toán tương quan món ăn.", body_style))
        
    story.append(Spacer(1, 15))
    
    # 4. Action recommendations
    story.append(Paragraph("4. Khuyến nghị Dinh dưỡng & Lối sống", h1_style))
    rec_text = """
    • <b>Khuyến nghị Uống nước:</b> Duy trì thói quen uống đều đặn 2.5 lít nước lọc ấm hàng ngày. Nước là dung môi hòa tan acid uric trong máu và thải qua nước tiểu.<br/>
    • <b>Chế độ dinh dưỡng:</b> Tránh ăn thịt chó, thịt bò, hải sản thân mềm. Tăng cường rau họ cải (bắp cải, cải thảo), trái cây có vị chua nhẹ (chứa Vitamin C giúp thải uric).<br/>
    • <b>Khi có đau khớp chân:</b> Tuyệt đối không tập thể thao nặng, chườm mát lên khớp sưng, gác chân cao hơn tim và theo dõi mức độ đau để kịp thời tư vấn bác sĩ gia đình.
    """
    story.append(Paragraph(rec_text, body_style))
    
    doc.build(story)
    output.seek(0)
    return output
